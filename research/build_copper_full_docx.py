#!/usr/bin/env python3
"""Build a full-length Word paper draft from COPPER_FULL_PAPER.md."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "research" / "COPPER_FULL_PAPER.md"
OUT = ROOT / "research" / "COPPER_CONFERENCE_DRAFT.docx"

INK = RGBColor(31, 45, 61)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(85, 94, 105)
HEADER_FILL = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(8.7)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.0

    for name, size, color, before, after in [
        ("Heading 1", 12.8, BLUE, 6, 2),
        ("Heading 2", 10.4, BLUE, 4, 1.5),
        ("Heading 3", 9.3, DARK_BLUE, 3, 1),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code = doc.styles.add_style("CodeBlock", 1)
    code.font.name = "Consolas"
    code.font.size = Pt(6.8)
    code.font.color.rgb = RGBColor(35, 39, 47)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(3)
    code.paragraph_format.line_spacing = 1.0

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(8.45)
        style.paragraph_format.space_after = Pt(1.5)
        style.paragraph_format.line_spacing = 1.0


def clean_inline(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("`", "")
    return text


def split_table_row(line: str) -> list[str]:
    cells = [clean_inline(cell.strip()) for cell in line.strip().strip("|").split("|")]
    return cells


def looks_like_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line))


def table_widths(headers: list[str]) -> list[int]:
    n = max(1, len(headers))
    total = 9600
    if n <= 3:
        return [int(total / n)] * n
    if n <= 7:
        first = 1600
        rest = int((total - first) / (n - 1))
        return [first] + [rest] * (n - 1)
    return [int(total / n)] * n


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if len(rows) < 2:
        return
    headers = rows[0]
    body = [row for row in rows[1:] if len(row) == len(headers)]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    font_size = Pt(5.5 if len(headers) > 7 else 6.1)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.size = font_size
    for row_values in body:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            run.font.size = font_size
    set_table_geometry(table, table_widths(headers))
    doc.add_paragraph()


def add_title(doc: Document, title_text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(clean_inline(title_text.lstrip("# ")))
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(15.5)
    run.font.color.rgb = INK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run(
        "Full research draft with ARM/AArch64 gem5 SE, Vivado RTL, and prior-art review"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(8.2)
    run.font.color.rgb = MUTED


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "COPPER research draft - public-knowledge novelty claim"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = MUTED


def flush_paragraph(doc: Document, paragraph_lines: list[str]) -> None:
    if not paragraph_lines:
        return
    text = clean_inline(" ".join(line.strip() for line in paragraph_lines).strip())
    if not text:
        return
    p = doc.add_paragraph()
    p.add_run(text)
    paragraph_lines.clear()


def build() -> None:
    lines = SRC.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style_document(doc)
    add_title(doc, lines[0])

    para: list[str] = []
    table_rows: list[list[str]] = []
    in_code = False
    code_lines: list[str] = []
    skip_artifact_summary = False
    in_references = False

    for raw in lines[1:]:
        line = raw.rstrip()
        heading_probe = re.match(r"^(#{2,4})\s+(.*)$", line)
        if heading_probe and heading_probe.group(2).strip() == "Artifact Summary":
            flush_paragraph(doc, para)
            if table_rows:
                add_table(doc, table_rows)
                table_rows.clear()
            skip_artifact_summary = True
            continue
        if heading_probe and heading_probe.group(1) == "##" and heading_probe.group(2).strip() == "References":
            skip_artifact_summary = False
            in_references = True
        elif heading_probe and heading_probe.group(1) == "##":
            in_references = False
        if skip_artifact_summary:
            continue

        if line.startswith("```"):
            if in_code:
                flush_paragraph(doc, para)
                p = doc.add_paragraph(style="CodeBlock")
                p.add_run("\n".join(code_lines))
                code_lines.clear()
                in_code = False
            else:
                flush_paragraph(doc, para)
                if table_rows:
                    add_table(doc, table_rows)
                    table_rows.clear()
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|") and "|" in line[1:]:
            flush_paragraph(doc, para)
            if looks_like_separator(line):
                continue
            table_rows.append(split_table_row(line))
            continue
        if table_rows:
            add_table(doc, table_rows)
            table_rows.clear()

        if not line.strip():
            flush_paragraph(doc, para)
            continue

        heading = re.match(r"^(#{2,4})\s+(.*)$", line)
        if heading:
            flush_paragraph(doc, para)
            level = min(len(heading.group(1)) - 1, 3)
            doc.add_heading(clean_inline(heading.group(2)), level=level)
            continue

        bullet = re.match(r"^\s*-\s+(.*)$", line)
        if bullet:
            flush_paragraph(doc, para)
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(clean_inline(bullet.group(1)))
            continue

        number = re.match(r"^\s*\d+\.\s+(.*)$", line)
        if number:
            flush_paragraph(doc, para)
            if in_references:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.22)
                p.paragraph_format.first_line_indent = Inches(-0.22)
                p.add_run(clean_inline(line))
            else:
                p = doc.add_paragraph(style="List Number")
                p.add_run(clean_inline(number.group(1)))
            continue

        para.append(line)

    flush_paragraph(doc, para)
    if table_rows:
        add_table(doc, table_rows)

    add_footer(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
