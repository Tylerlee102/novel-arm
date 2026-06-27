#!/usr/bin/env python3
"""Build a focused COPPER/SCOOP conference-style DOCX draft."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "research" / "COPPER_SCOOP_CONFERENCE_DRAFT.docx"

INK = RGBColor(31, 45, 61)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(91, 103, 116)
HEADER_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top: int = 60, bottom: int = 60, start: int = 90, end: int = 90) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.88)
    section.bottom_margin = Inches(0.88)
    section.left_margin = Inches(0.90)
    section.right_margin = Inches(0.90)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.55)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(4.6)
    normal.paragraph_format.line_spacing = 1.06

    for name, size, color, before, after in [
        ("Heading 1", 15.25, BLUE, 11, 5.5),
        ("Heading 2", 12.55, BLUE, 8.5, 4.25),
        ("Heading 3", 11.55, DARK_BLUE, 5.5, 3.25),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(10.0)
        style.paragraph_format.space_after = Pt(4.1)
        style.paragraph_format.line_spacing = 1.06
        style.paragraph_format.left_indent = Inches(0.30)
        style.paragraph_format.first_line_indent = Inches(-0.16)

    callout = doc.styles.add_style("Callout", 1)
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.0)
    callout.font.color.rgb = INK
    callout.paragraph_format.left_indent = Inches(0.10)
    callout.paragraph_format.right_indent = Inches(0.10)
    callout.paragraph_format.space_before = Pt(4.25)
    callout.paragraph_format.space_after = Pt(5.0)
    callout.paragraph_format.line_spacing = 1.06


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    run = title.add_run(
        "COPPER and SCOOP: Committed Pointer-Provenance Authority for Safe Data-Dependent Prefetching"
    )
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(15.2)
    run.font.color.rgb = INK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(6)
    run = subtitle.add_run(
        "Focused architecture/security draft - AArch64 full-system gem5, public workloads, adversarial oracles, and Vivado RTL evidence"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(8.1)
    run.font.color.rgb = MUTED


def add_footer(doc: Document) -> None:
    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "COPPER/SCOOP focused research draft - public-knowledge novelty claim"
    footer.runs[0].font.size = Pt(7.3)
    footer.runs[0].font.color.rgb = MUTED


def paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.add_run(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    font_size = Pt(6.45 if len(headers) > 5 else 7.05)
    for idx, header in enumerate(headers):
        cell = tbl.rows[0].cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        run.bold = True
        run.font.size = font_size
    for row_values in rows:
        row = tbl.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            if idx > 0 and len(value) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            run.font.size = font_size
    set_table_geometry(tbl, widths)
    doc.add_paragraph()


def callout(doc: Document, text: str) -> None:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, CALLOUT_FILL)
    set_cell_margins(cell, 90, 90, 140, 140)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(8.3)
    set_table_geometry(tbl, [9300])
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("Abstract", level=1)
    paragraph(
        doc,
        "Data-memory-dependent prefetchers (DMPs) can accelerate irregular pointer-heavy code, but Augury and GoFetch show that pointer-looking data can also trigger secret-dependent cache activity even when software never uses that data as an address. COPPER changes the DMP authority model: a data-derived prefetch may issue only when committed execution has proven the exact source word as a pointer source, the source remains clean, the address-space/protection token matches, and recursive cross-page targets have committed target-line witnesses. This paper adds SCOOP, Slack-only COPPER Companion Prefetching, which runs COPPER as a content-derived companion beside a conventional primary prefetcher; the primary has strict issue priority and COPPER issues only in slack cycles. In AArch64 full-system gem5, unsafe DMP emits 32,760 extra prefetches when a secret bit changes data words from high-bit junk to heap addresses. COPPER and SCOOP reduce the unauthorized scan-phase allowed-candidate delta to zero; a split scan/probe audit shows unsafe DMP leaks before the observer probe, while COPPER/SCOOP block the scan-phase candidates. A 12-point SQLite/Lua/Duktape/yyjson plus JSON+SQLite and cache-service app matrix shows SPP is the best conventional baseline, while SCOOP stays within 0.360 percentage points worst-case and preserves COPPER's authority checks and zero translation faults; a 15-point public-engine seed portfolio cuts naive CTLW misses by 90.706%. A 22-point gem5-counter scorecard gives standalone COPPER an 18.8% lower base-weighted mean pollution proxy than naive DMP, with 18.1%-20.6% lower results across six checked weight scenarios, and quantifies SCOOP's incremental traffic over SPP. Vivado XSim and bounded checkers verify SCOOP arbitration plus OoO-LSQ and TLB/coherence contracts; the TLB/coherence issue filter passes 27 directed plus 10,000 randomized XSim checks and synthesizes with WNS +6.898 ns. To the best of public knowledge, COPPER/SCOOP is the first public DMP defense to make committed pointer provenance the authority for recursive content-derived prefetching while coexisting with conventional address-correlation prefetchers.",
    )

    doc.add_heading("1. Problem and Thesis", level=1)
    paragraph(
        doc,
        "A DMP uses loaded memory contents as future addresses. That is attractive for linked structures, hash tables, runtimes, and graph-like code, but it changes constant-time assumptions: a secret value that merely looks like a pointer can cause a cache fill. Disabling all such prefetching is safe but blunt. COPPER asks a narrower architectural question: can a DMP retain useful pointer runahead while refusing to dereference arbitrary data?",
    )
    callout(
        doc,
        "COPPER invariant: address-shapedness is not authority. A DMP may dereference a source word only after committed architectural execution has already used that same word/value as a pointer source under the current address-space authority, and the proof has not been invalidated.",
    )
    bullets(
        doc,
        [
            "Security goal: no content-derived prefetch from data-at-rest, stale, cross-domain, or permission-invalid source words.",
            "Non-goal: hiding legitimate architectural memory accesses or proving all software constant-time.",
            "Positioning: COPPER is a safe authority layer for content-derived prefetching; SCOOP composes it with conventional address-stream prefetchers.",
        ],
    )

    doc.add_heading("2. Threat Model and Claim", level=1)
    paragraph(
        doc,
        "The attacker can observe cache timing and can influence or observe data values that may be loaded by victim code. The victim may follow constant-time software rules at the ISA level: secret data does not control branches or architectural memory addresses. A vulnerable DMP breaks that contract by issuing prefetches from loaded data values that merely resemble addresses. COPPER targets exactly this microarchitectural gap. It does not claim to stop cache leakage from legitimate architectural accesses, speculative execution outside the DMP, or software that actually computes a secret-dependent address.",
    )
    paragraph(
        doc,
        "The paper's central security claim is differential. If two executions differ only in whether ordinary loaded data words contain pointer-shaped secret values, the content-derived COPPER/SCOOP lane should not create a corresponding change in unauthorized prefetch issue. This is stronger than checking that a run has few prefetches, because a small but secret-correlated delta can still be a side channel. The new oracle experiments therefore compare `secret=1` against `secret=0`, report issued-prefetch and allowed-candidate deltas, and add an observer phase that tests whether those prefetches warm measurable cache state.",
    )
    paragraph(
        doc,
        "The intended deployment boundary is an AArch64-style core or CPU cluster with ordinary virtual memory, context tokens, and coherent SoC events. The implementation avoids relying on proprietary Arm microarchitecture. PASB can be modeled as an ASID/VMID/security-state token, CTLW as a committed virtual-to-physical line witness, and CS-SARI as a candidate-specific interface from coherence, DMA, remap, and TLBI events into DMP authority revocation.",
    )

    doc.add_heading("3. COPPER Mechanism", level=1)
    paragraph(
        doc,
        "COPPER tracks positive prefetch authority, not secrecy. A proof is created only when committed demand execution uses a loaded source word as an address source and translation/permission succeeds. Stores, fills, invalidations, DMA/coherent writes, context-token changes, and witness revocations clear or block proof. The implementation evaluated here combines several named structures.",
    )
    table(
        doc,
        ["Mechanism", "Role"],
        [
            ["RCP", "Recursive Carried-Provenance; prefetched lines carry identity/context, not new authority."],
            ["CEPF", "Commit-Epoch Provenance Filter; prevents stale in-flight source tags from recreating proof after overwrite."],
            ["PASB", "Provenance Address-Space Binding; keys proofs to address-space/protection tokens."],
            ["CTLW", "Committed Target-Line Witnessing; exact demand-observed target-line translation witness for recursive cross-page issue."],
            ["CLPD", "Compressed Line-Provenance Directory; stores retained line proof with per-word masks and epochs."],
            ["PEB", "Provenance Epoch Boundary; cancels pre-boundary proof authority without sweeping the whole directory."],
            ["CS-SARI", "Conflict-scoped SoC revocation; holds only candidates conflicting with pending DMA/coherence/remap/TLBI authority changes."],
        ],
        [1180, 8120],
    )

    paragraph(
        doc,
        "A subtle point is that proof is positive authority, not permission by absence of danger. The prefetcher cannot infer that a word is safe because it has a canonical address shape, because the target translation succeeds, or because the target line is present. Those facts are useful only after source authority exists. This is why source provenance, address-space binding, and target-line witnessing must be composed: source proof answers whether the word may be used as a DMP source, PASB answers whether the proof belongs to the current protection context, and CTLW answers whether a recursive cross-page target line has a committed translation witness.",
    )

    doc.add_heading("4. SCOOP Hybrid", level=1)
    paragraph(
        doc,
        "Standalone COPPER is intentionally conservative: first-use DMP prefetching is lost until committed proof exists. Reviewers also care that conventional prefetchers such as SPP/DCPT/AMPM often win raw timing on public workloads. SCOOP addresses that objection by making COPPER a slack-only companion lane rather than a replacement.",
    )
    callout(
        doc,
        "SCOOP invariant: if the conventional primary lane has a ready prefetch, the COPPER companion lane must not issue. COPPER may issue only in primary slack cycles.",
    )
    paragraph(
        doc,
        "This invariant is small but important: it gives the conventional prefetcher full priority for address-correlation performance while preserving a separated content-derived authority path. A bounded checker passes SCOOP to depth 10 and finds immediate counterexamples for companion-first and round-robin weakened variants. A synthesizable arbiter passes Vivado XSim with 6 directed plus 10,000 randomized cases, `companion_blocks=2360`, and `errors=0`.",
    )

    doc.add_heading("5. Evaluation Methodology", level=1)
    paragraph(
        doc,
        "The evaluation has three layers. First, trace and bounded-state models isolate the invariant and generate counterexamples for weakened rules. Second, RTL checks validate the gate-level and arbiter-level structures under directed and randomized tests. Third, gem5 full-system runs boot Linux and execute native AArch64 binaries with the selected prefetcher attached at the L1D cache hierarchy. The full-system runs are deliberately separated into adversarial security oracles and public workload controls.",
    )
    table(
        doc,
        ["Layer", "What it measures", "Why reviewers should care"],
        [
            ["Adversarial oracles", "Secret-dependent prefetch deltas, L1D miss deltas, scan/probe phase splits", "Directly tests the Augury/GoFetch-style failure mode."],
            ["Public applications", "SQLite, Lua, Duktape, yyjson, JSON+SQLite, Olden, GAPBS, heap layouts", "Checks correctness, timing, and conventional-baseline pressure on non-toy code."],
            ["RTL/checkers", "SCOOP arbitration, CLPD/PEB/authority gates, state-space counterexamples", "Shows the mechanism is not just a simulator flag."],
            ["Cost proxies", "LUT/FF/BRAM/timing, proof storage, blocked/allowed/translation counters", "Bounds implementation plausibility and remaining integration risk."],
        ],
        [1650, 3550, 4100],
    )
    paragraph(
        doc,
        "All gem5 full-system rows used in the new security portfolio complete with `rc=0` and checksum validation. The paper reports conventional prefetchers where they are stronger: DCPT/SPP/AMPM often beat standalone COPPER on raw timing. That does not invalidate the contribution because those baselines do not provide a safe authority rule for content-derived DMP candidates. It does force the paper to avoid a universal performance claim.",
    )

    doc.add_heading("6. Security Evidence", level=1)
    table(
        doc,
        ["Evidence", "Unsafe DMP signal", "COPPER", "SCOOP"],
        [
            ["Fake-only ROI", "28,685 PF from 131,094 fake observations", "PEB allows 0, blocks 131,066", "Companion allows 0, blocks 191,674"],
            ["Secret traffic oracle", "PF delta 32,760; allowed delta 32,760", "Allowed delta 0; blocked delta 32,760", "Allowed delta 0; blocked delta 64,143"],
            ["Cold-cache observer", "L1D miss delta -14; timing shift -4.906 pp", "L1D delta 0; allowed delta 1", "L1D delta 1; allowed delta 0"],
            ["3-seed observer", "Allowed deltas 63..65; L1D deltas -14..-9", "L1D delta 0 all seeds", "Allowed delta set 0"],
            ["Split scan/probe", "Scan PF delta 64; allowed delta 66", "Scan allowed delta 0; blocked delta 64", "Scan allowed delta 0; blocked delta 64"],
        ],
        [1900, 2650, 2350, 2400],
    )
    paragraph(
        doc,
        "The split audit is the cleanest security result. It inserts a stats dump/reset between the secret data scan and the later target-probe phase. Unsafe DMP has a scan-phase allowed delta of 66 before any architectural target probe. COPPER and SCOOP both have scan-phase allowed delta 0 and blocked delta 64. This supports the precise differential claim: secret-shaped data changes unsafe DMP traffic, but not COPPER/SCOOP content-derived issue authority.",
    )
    paragraph(
        doc,
        "The cold-cache observer oracle adds a second check: prefetch traffic must become cache state to be exploitable. The workload first evicts cache state, then scans secret-dependent data words, then probes the target lines. For unsafe DMP, `secret=1` reduces target-probe L1D demand misses by 14 and shifts policy timing by -4.906 percentage points relative to `secret=0`. Across three address permutations, unsafe DMP always has positive allowed deltas (63..65) and fewer L1D misses for `secret=1` (-14..-9). SCOOP's companion allowed delta remains 0 for all three seeds.",
    )
    paragraph(
        doc,
        "The traffic-only oracle and observer oracle intentionally answer different questions. The traffic-only oracle keeps the program from touching target lines after the scan, so useful-prefetch counters are not the main signal; the issue/allowed deltas are. The observer oracle then asks whether the unsafe traffic becomes measurable cache state. The split audit resolves the main ambiguity of the observer setup: once the program begins probing target lines, some later prefetch behavior can be a legitimate response to architectural dereference. By dumping scan statistics before the probe, the audit shows the unauthorized leak is already present for unsafe DMP and absent for COPPER/SCOOP.",
    )
    table(
        doc,
        ["Ablation", "Observed failure", "Lesson"],
        [
            ["Naive DMP", "Secret-dependent PF and allowed deltas; observer L1D miss reduction", "Address-shaped data is not safe authority."],
            ["COPPER without boundary", "Warm proof state can leak into fake-only measurement", "A provenance epoch boundary is needed around security domains/ROIs."],
            ["Companion-first/round-robin SCOOP variants", "Bounded checker finds immediate primary-ready violations", "Strict primary priority is the hybrid invariant."],
            ["Source-only provenance", "Earlier trace sweeps permit stale rewritten-slot issues", "Value/epoch and clean-since-proof lifecycle matter."],
            ["No CTLW/terminal rule", "PASB-only full-system runs leave recursive translation hazards", "Recursive cross-page issue needs exact target-line witnesses."],
        ],
        [2000, 3300, 4000],
    )

    doc.add_heading("7. Full-System Workloads", level=1)
    paragraph(
        doc,
        "The evaluation now includes public application-style engines and public pointer-intensive benchmarks running as native AArch64 Linux binaries under gem5 full-system. These are not presented as universal performance wins; they are used to test whether the authority mechanism composes with real software and conventional prefetch baselines.",
    )
    table(
        doc,
        ["Workload", "Best conventional", "SCOOP timing", "SCOOP blocked", "Faults", "Readout"],
        [
            ["SQLite medium", "SPP -3.623%", "-3.617%", "72,512", "0", "SPP-class timing; CTLW 1,778"],
            ["SQLite stress", "SPP -2.587%", "-2.549%", "175,838", "0", "COPPER CTLW 2,543 vs naive 43,226"],
            ["Lua medium", "SPP -29.532%", "-29.240%", "247,148", "0", "COPPER faster than naive; CTLW -91.3%"],
            ["Lua stress", "SPP -31.392%", "-31.120%", "655,685", "0", "COPPER faster than naive; CTLW -76.9%"],
            ["Duktape medium", "SPP -6.732%", "-6.950%", "163,077", "0", "Public JS engine; checksum stable"],
            ["Duktape stress", "SPP -8.385%", "-8.745%", "301,435", "0", "Slack hybrid exceeds SPP timing in this run"],
            ["yyjson medium", "SPP -18.351%", "-18.342%", "19,465", "0", "Public JSON parser; CTLW -98.9%"],
            ["yyjson stress", "SPP -22.097%", "-22.186%", "41,408", "0", "Stress JSON parse/traverse; CTLW -98.9%"],
            ["JSON+SQLite medium", "SPP -4.497%", "-4.523%", "78,860", "0", "Public parser+database composition; CTLW -95.0%"],
            ["JSON+SQLite stress", "SPP -3.588%", "-3.623%", "222,066", "0", "Larger parser+database composition; CTLW -91.4%"],
            ["Cache-service small", "SPP -13.440%", "-13.406%", "7,216", "0", "Hash/LRU service pattern; CTLW -99.5%"],
            ["Cache-service medium", "SPP -13.115%", "-13.086%", "9,967", "0", "Larger hash/LRU service pattern; CTLW -99.4%"],
            ["Olden random small", "DCPT -5.742%", "-3.192%", "1,211,472", "0", "Conventional wins raw timing; COPPER blocks content-derived risk"],
            ["Heap fake-only", "SPP -53.512%", "-53.487%", "191,674", "0", "All fake content candidates blocked"],
        ],
        [1600, 1500, 1300, 1350, 760, 2790],
    )
    paragraph(
        doc,
        "The strongest honest performance statement is not that COPPER beats every prefetcher. It does not. On the 12-point app matrix, standalone COPPER is faster than or equal to naive DMP on 5/12 and has fewer L1D demand misses on 8/12; after adding stride, DCPT, AMPM, and SPP to every point, SPP is the best conventional baseline on all 12. The stronger claim is that SCOOP preserves conventional-prefetcher performance where the primary lane works: mean SPP timing is -13.112% versus no prefetch, mean SPP+COPPER slack timing is -13.116%, the average signed gap is -0.004 percentage points, and the worst absolute gap is 0.360 points. A 15-point medium/stress SQLite/Lua/Duktape seed portfolio adds layout stability evidence: COPPER cuts aggregate naive CTLW misses by 90.706%, SPP+COPPER slack cuts them by 91.505%, and the worst slack-SPP gap is 0.760 points. COPPER supplies a separately verified authority gate for content-derived DMP candidates.",
    )
    paragraph(
        doc,
        "A new gem5-counter pressure scorecard separates speed from side effects. Using a transparent weighted proxy over bus bytes, DRAM reads, L2 replacements, and L1D replacements, standalone COPPER scores 0.879% versus 1.083% for naive DMP under the base weights, an 18.8% lower mean proxy pollution score; across six checked weight scenarios, the reduction ranges from 18.1% to 20.6%. SCOOP inherits SPP's high-traffic profile, but the incremental cost is now quantified: it is within 0.5% runtime of SPP on all 22 points and adds 0.093 bus-byte-delta points and 0.222 pressure-score points over SPP on average. This is not calibrated power.",
    )
    paragraph(
        doc,
        "SQLite, Lua, Duktape, yyjson, a composed JSON+SQLite workload, and cache-service hash/LRU scale points were added specifically to reduce the risk that the paper was only a synthetic pointer benchmark. Each is compiled into a native AArch64 Linux workload and run under the same full-system path, with medium/stress scale points for the single-engine families, repeated layouts for the first three engines, bounded medium/stress service-composition points, and small/medium cache-service hash/LRU scale points. They do not prove broad real-world performance, but they show that the mechanism survives realistic allocator, table, object, interpreter, database, JSON parse/traverse, and hash/LRU update behaviors with stable checksums, zero fill-origin translation faults, and 91.1% aggregate CTLW reduction versus naive DMP across the 12-point app matrix. The full conventional baseline matrix prevents overclaiming: ordinary address-stream prefetching still wins raw speed, while the companion path retains SPP-class timing but inherits SPP-sized traffic. Standalone COPPER averages +0.754% bus bytes and should be framed as the low-overhead authority path.",
    )
    paragraph(
        doc,
        "The heap fake-only control is an especially important negative control. The ROI contains pointer-shaped words but no architectural pointer traversal. Naive DMP turns those values into 28,685 content-derived prefetches and severe slowdown. CLPD-64K with PEB blocks 131,066 of 131,066 observations and issues zero prefetches. SCOOP preserves the fast SPP stream behavior while the companion lane blocks 191,674 unproven content-derived candidates. This is the experiment that prevents the paper from claiming safety only on benign pointer-heavy code.",
    )
    doc.add_heading("8. Result Interpretation", level=1)
    paragraph(
        doc,
        "The results create a deliberately asymmetric story. On security, unsafe DMP gives large, repeatable secret-dependent signals and COPPER/SCOOP remove the unauthorized scan-phase issue. On performance, conventional address-correlation prefetchers often remain best. A strong submission should not hide that asymmetry. It should argue that modern cores already deploy multiple prefetchers, and that a safe content-derived lane should be judged by whether it admits DMP-like candidates without violating the software address-generation contract.",
    )
    paragraph(
        doc,
        "SCOOP is therefore the better paper mechanism than standalone COPPER alone. Standalone COPPER proves that committed provenance can gate content-derived issue, but its conservative first-use and exact target-line witness rules give up some raw opportunity. SCOOP keeps the conventional prefetcher in front and turns COPPER into a filtered opportunistic companion. The application table shows this clearly: SQLite, Lua, Duktape, yyjson, and JSON+SQLite runs retain SPP-class timing, while the companion continues to block tens to hundreds of thousands of content-derived candidates and records zero fill-origin translation faults.",
    )
    paragraph(
        doc,
        "The security oracles also justify the slack-only design. If COPPER were simply mixed into a round-robin multi-prefetcher, it could issue while the primary lane had useful work ready, creating both performance interference and a harder-to-explain arbitration policy. SCOOP instead has one invariant reviewers can check: primary-ready implies no companion issue. The bounded checker and RTL arbiter make that invariant executable. The mechanism is simple enough to explain, but the measured behavior is not a trivial block combination: conventional address prediction and committed-provenance DMP authority coexist without granting address-shaped data a prefetch right.",
    )
    paragraph(
        doc,
        "The remaining performance question is where safe content-derived prefetching has enough natural demand to matter beyond controlled heaps and pointer structures. Public engines exercise realistic allocators and object/table behavior, but a top-tier evaluation would still want SPEC-like C/C++ applications, language-runtime workloads at larger scale, and crypto-adjacent code where DMP leakage risk is concrete. The present evidence is enough to make the architecture/security claim credible; it is not enough to claim a universal speedup.",
    )

    doc.add_heading("9. RTL, Models, and Cost", level=1)
    table(
        doc,
        ["Artifact", "Result"],
        [
            ["SCOOP arbiter RTL", "Vivado XSim PASS; 6 directed + 10,000 randomized cases; errors=0."],
            ["SCOOP state checker", "PASS to depth 10; companion-first and round-robin fail as expected."],
            ["CLPD SRAM directory", "64K-entry config synthesizes on Artix-7 200T: 629 LUTs, 156 FFs, 260 BRAM tiles; route WNS 0.362 ns."],
            ["PEB RTL", "346 LUTs, 147 FFs, no BRAM/DSP, WNS 3.782 ns at 10 ns."],
            ["Authority checkers", "CEPF/PASB/CTLW, CLPD, OoO-LSQ, and TLB/coherence checks pass; weakened variants fail."],
            ["TLB/coherence filter", "XSim PASS: 27 directed + 10,000 randomized; 332 LUTs, 167 FFs, WNS +6.898 ns."],
            ["CS-SARI", "Composition checker passes 7,555 states; no-hold control produces stale-authority counterexamples."],
        ],
        [2250, 7050],
    )
    paragraph(
        doc,
        "The proof path still needs production backend integration: a real load-store queue must carry source-word/value tags to committed dependent memory operations, and a production SoC must connect DMA, CHI/ACE/AXI coherence, remap, and TLBI events to proof revocation. The bounded OoO-LSQ checker makes the backend obligation executable. A new TLB/coherence contract explores 39,098 full-contract states, and its RTL issue filter passes 27 directed plus 10,000 randomized XSim checks while synthesizing to 332 LUTs / 167 FFs with WNS +6.898 ns. These are contract evidence, not a production core.",
    )
    paragraph(
        doc,
        "The SCOOP arbiter is intentionally tiny, because its value is the invariant rather than raw area. The larger cost comes from COPPER metadata: source proof storage, target witnesses, and revocation bookkeeping. CLPD is the main storage response: it compresses retained proof by cache line with per-word masks and epochs, and the 64K-entry SRAM directory has an out-of-context route result on Artix-7. PEB is a small boundary hook, not a novel epoch concept by itself. The novelty is binding DMP proof authority to epochs so pre-boundary authority cannot survive into post-boundary measurement windows.",
    )

    doc.add_heading("10. Prior Art and Novelty Risk", level=1)
    table(
        doc,
        ["Prior art", "Overlap", "Difference"],
        [
            ["Augury / GoFetch", "DMP side-channel threat", "Attacks/reverse engineering; no committed source-word authority defense."],
            ["SplittingSecrets", "DMP defense", "Compiler transforms secrets; COPPER changes hardware authority."],
            ["PreFence / DIT / DOIT", "Disable or schedule around prefetch risk", "Coarse policy; COPPER is per-candidate and keeps safe activity."],
            ["Pointer-chase, indirect, ICP", "Irregular prefetch performance", "Do not require clean committed pointer-source proof for issue."],
            ["CHERI / MTE / taint", "Metadata and provenance", "Architectural safety or information-flow tracking, not DMP-specific positive authority."],
        ],
        [2100, 2100, 5100],
    )
    paragraph(
        doc,
        "The novelty claim must stay scoped: to the best of public knowledge, no public mechanism was found that uses clean committed pointer-source provenance plus address-space and target-line witnesses as the authority for recursive DMP issue, or that composes that authority as a strict-priority slack companion to conventional prefetching. Commercial cores may have unpublished DMP safety rules.",
    )
    paragraph(
        doc,
        "A skeptical reviewer may call COPPER a combination of known blocks: metadata, prefetching, epochs, and translation checks. The response is that the paper should not claim novelty for those blocks. The contribution is the named authority rule and its measurable behavior: source-word proof is created only by committed pointer use, recursive issue cannot bootstrap from prefetched data alone, and the slack-only companion can preserve primary prefetcher priority while remaining differentially silent under secret-shaped data. That behavior is what the oracles measure.",
    )

    doc.add_heading("11. Limitations", level=1)
    bullets(
        doc,
        [
            "First-use DMP prefetching is lost by design; safe proof seeding is future work.",
            "Public applications are broader than before and now include SQLite/Lua/Duktape/yyjson medium/stress runs, bounded JSON+SQLite medium/stress service-composition points, bounded cache-service hash/LRU scale points, and repeated SQLite/Lua/Duktape layouts, but SPEC-like, production database, runtime, and crypto-adjacent workloads remain important.",
            "Energy, L2 pollution, NoC pressure, and memory-command overhead now have a gem5-counter proxy scorecard, but calibrated power modeling remains future work.",
            "The backend proof path and full AMBA/CHI revocation integration need production-style RTL or formal models beyond the bounded OoO-LSQ and TLB/coherence contracts.",
            "SCOOP is a mechanism and artifact result, not a guarantee of top-conference acceptance.",
        ],
    )
    paragraph(
        doc,
        "The paper also needs careful language around acceptance probability. The artifact is substantially stronger after the full-system oracles, public engine workloads including yyjson and JSON+SQLite, repeated SQLite/Lua/Duktape medium/stress layouts, SCOOP hybrid, Vivado arbiter check, OoO-LSQ and TLB/coherence contract checkers, and seed sweep. That still does not make a top-tier PhD conference acceptance guaranteed. Top-tier reviewers may require SPEC-like applications, production-quality backend and memory-system proof paths, and stronger energy/pollution analysis. The current honest verdict is focused-conference plausible and workshop-strong, with a credible path to a larger submission.",
    )

    doc.add_heading("12. Reproducibility Package", level=1)
    paragraph(
        doc,
        "The artifact is organized so a reviewer can audit the claim from multiple directions. The gem5 implementation lives in the local gem5 tree as COPPER and CopperCompanion prefetcher models. The AArch64 full-system runner exposes policy switches for `none`, `naive`, `copper_clpd64k_peb`, `spp`, and `spp_copper_slack`. Public workload build scripts compile SQLite, Lua, Duktape, yyjson, JSON+SQLite, Olden, GAPBS-style, and adversarial oracle binaries against the installed AArch64 sysroot. The security summaries are generated from gem5 `stats.txt` and terminal output rather than manually entered tables.",
    )
    table(
        doc,
        ["Artifact", "Purpose"],
        [
            ["COPPER_SECURITY_EVIDENCE_PORTFOLIO_20260616.md", "Consolidates fake-only, traffic-oracle, observer, seed-sweep, split-audit, checker, and RTL evidence."],
            ["DMP_ORACLE_I8192_P4_SUMMARY.md", "Traffic-only secret-bit oracle: unsafe PF delta 32,760; COPPER/SCOOP allowed deltas 0."],
            ["DMP_ORACLE_OBSERVER_SEED_SWEEP.md", "Three-seed cold-cache observer oracle: unsafe L1D miss deltas -14..-9; SCOOP allowed delta set 0."],
            ["DMP_ORACLE_I512_P4_PROBE1_EVICT512_SPLIT_SUMMARY.md", "Separates unauthorized scan from legitimate target probe; COPPER/SCOOP scan allowed deltas 0."],
            ["COPPER_APPLICATION_WORKLOAD_PORTFOLIO_20260616.md", "SQLite/Lua/Duktape/yyjson/JSON+SQLite public workload summary with SPP/SCOOP comparison."],
            ["COPPER_PUBLIC_APP_REPEATED_SEED_PORTFOLIO_20260617.md", "Fifteen medium/stress SQLite/Lua/Duktape seeds: COPPER CTLW -90.706%; SCOOP CTLW -91.505%."],
            ["COPPER_ENERGY_POLLUTION_SCORECARD_20260617.md", "Traffic/pollution proxy score: COPPER 18.8% lower than naive under base weights, 18.1%-20.6% lower across checked weights; SCOOP incremental cost over SPP quantified."],
            ["COPPER_OOO_LSQ_PROOF_CONTRACT.md", "Bounded OoO-LSQ retirement proof contract; full rule passes and weakened variants fail."],
            ["COPPER_TLB_COHERENCE_CONTRACT.md", "Bounded target-witness/remap/TLBI/permission contract; full rule passes and weakened variants fail."],
            ["COPPER_TLB_COHERENCE_AUTHORITY_FILTER_RTL_SUMMARY.md", "Synthesizable issue-side TLB/coherence filter; XSim and timing pass."],
            ["COPPER_SLACK_COMPANION_HYBRID_20260615.md", "SCOOP mechanism summary, workload table, bounded checker, and Vivado XSim arbiter result."],
        ],
        [3000, 6300],
    )
    paragraph(
        doc,
        "A reviewer can rerun the focused oracles without rebuilding the whole paper: build the AArch64 oracle binary, run the full-system policy matrix for both secrets, and regenerate the summaries. The split audit is especially useful for debugging because it produces two stats blocks per run, one for the secret scan and one for the observer probe. The Vivado arbiter test is separate and quick compared with gem5: it checks the SCOOP priority invariant under directed and randomized ready/stall patterns. These pieces are intentionally smaller than a full production core so the artifact remains inspectable.",
    )

    doc.add_heading("13. Conclusion", level=1)
    paragraph(
        doc,
        "COPPER reframes data-dependent prefetching around authority: a DMP should not dereference data because it looks like an address, but because committed execution has already proven the exact source word/value under the current context. SCOOP makes this authority path practical in a modern prefetch stack by giving conventional address-correlation prefetching strict priority and letting COPPER issue only in slack cycles. The local evidence now includes AArch64 full-system adversarial oracles, public application workloads including JSON+SQLite composition, multi-seed observer and public-engine layout runs, split scan/probe localization, bounded OoO-LSQ and TLB/coherence contracts, a matching TLB/coherence RTL filter, and Vivado RTL. The result is a serious conference-style candidate. The remaining work is not cosmetic: broaden application evidence, formalize the production integration boundary, and keep the novelty claim public-knowledge scoped.",
    )

    doc.add_heading("References", level=1)
    bullets(
        doc,
        [
            "Augury: Using Data Memory-Dependent Prefetchers to Leak Data at Rest. https://www.prefetchers.info/augury.pdf",
            "GoFetch: Breaking Constant-Time Cryptographic Implementations Using Data Memory-Dependent Prefetchers. https://gofetch.fail/",
            "SplittingSecrets. https://arxiv.org/abs/2601.12270",
            "PreFence. https://arxiv.org/abs/2410.00452",
            "Intel Data Dependent Prefetcher guidance. https://www.intel.com/content/www/us/en/developer/articles/technical/software-security-guidance/technical-documentation/data-dependent-prefetcher.html",
            "GAP Benchmark Suite. https://github.com/sbeamer/gapbs",
            "gem5. https://www.gem5.org/",
            "SQLite amalgamation, Lua 5.4.8, Duktape 2.7.0, and yyjson 0.12.0 public source distributions used for application workloads, including the composed JSON+SQLite service-style run.",
        ],
    )

    add_footer(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
