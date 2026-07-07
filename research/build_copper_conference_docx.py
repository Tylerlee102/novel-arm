#!/usr/bin/env python3
"""Build a compact Word draft for the COPPER paper package."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "research" / "COPPER_CONFERENCE_DRAFT.docx"


INK = RGBColor(31, 45, 61)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 99, 110)
HEADER_FILL = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_margins(section, margin: float = 0.72) -> None:
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)


def style_document(doc: Document) -> None:
    set_margins(doc.sections[0])
    styles = doc.styles

    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    for name, size, color, before, after in [
        ("Heading 1", 14, BLUE, 11, 5),
        ("Heading 2", 11.5, BLUE, 8, 4),
        ("Heading 3", 10.5, DARK_BLUE, 6, 3),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(9.3)
        style.paragraph_format.space_after = Pt(2)
        style.paragraph_format.left_indent = Inches(0.24)
        style.paragraph_format.first_line_indent = Inches(-0.14)


def add_run(paragraph, text: str, bold: bool = False, italic: bool = False):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    return run


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run(
        "COPPER: Committed Pointer-Provenance Prefetching for Safe "
        "Data-Memory-Dependent Prefetchers"
    )
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(17)
    run.font.color.rgb = INK

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    run = subtitle.add_run(
        "ARM/AArch64-style CPU subsystem research draft - updated with "
        "ARM gem5 SE/full-system CTLW timing ROI, Vivado RTL, and public prior-art review"
    )
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = MUTED


def add_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    p.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              widths_dxa: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, HEADER_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(7.7)
    for row_values in rows:
        row = table.add_row()
        for idx, text in enumerate(row_values):
            cell = row.cells[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(text)
            run.font.size = Pt(7.4)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph()


def add_footer(doc: Document) -> None:
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.text = "COPPER research draft - generated from local artifacts"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = MUTED


def build() -> None:
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("Abstract", level=1)
    add_paragraph(
        doc,
        "Data-memory-dependent prefetchers (DMPs) improve irregular "
        "pointer-heavy workloads by using memory contents as future addresses, "
        "but Augury and GoFetch show that this can violate constant-time "
        "assumptions: data that merely resembles a pointer may trigger "
        "secret-dependent cache activity. COPPER changes the DMP authority "
        "model. A DMP may dereference a memory-derived value only when the "
        "exact source word has committed pointer provenance, remains clean "
        "since proof, matches the protection/address-space context, and has a committed target-line "
        "translation witness when recursive cross-page issue is needed. The refined mechanism adds Recursive "
        "Carried-Provenance (RCP) runahead and a Committed Page-Translation "
        "Queue (CPTQ), plus Provenance Address-Space Binding (PASB), Committed Target-Line Witnessing (CTLW), and a Compressed Line-Provenance Directory (CLPD), allowing deeper pointer prefetching without allowing a "
        "prefetched line to create authority by itself. Trace simulation blocks "
        "all modeled unsafe DMP dereferences while retaining pointer-chain "
        "speedup. A GAPBS-backed topology trace over five generated Kronecker graphs shows COPPER-epoch at 1.770x and CLPD at 1.896x with zero data-at-rest, unproven-edge, or stale-slot prefetches; CLPD recovers a 2.115x g12 edge-scan speedup with 8,192 line entries where an edge-exact ledger needs 131,072 entries, and a storage proxy shows 54.00 KiB versus 1696.00 KiB at those measured points. An expanded GAPBS-style kernel sensitivity sweep over 4,320 graph/kernel/table/cache/lookahead runs keeps COPPER unsafe modeled prefetches at zero, while naive DMP produces 81,605,320 unsafe modeled prefetches and source-only provenance still produces 284,488. Vivado synthesis of the core gate meets a 10 ns Artix-7 "
        "constraint, a full-authority CEPF/PASB/CTLW RTL gate passes 12 "
        "directed plus 5,000 randomized XSIM tests, and the CLPD directory RTL "
        "passes 14 directed plus 5,000 randomized scoreboard tests; a fresh authority-chain regression passes 10 XSIM scripts with 0 failures. ARM/AArch64 gem5 syscall-emulation runs show recursive "
        "COPPER improves ARM32 page-permuted pointer chains by 6.76-6.78% and "
        "random chains by 5.59-5.66%; direct AArch64 ELF runs show 6.77% and "
        "5.61% on the corresponding full-list shapes. CS-SARI, a conflict-scoped "
        "revocation hold, preserves zero modeled stale-authority issues while "
        "recovering 269,879 authorized GAPBS-topology candidate opportunities "
        "that global SARI would hold; a 20-configuration queue-depth/conflict "
        "sweep keeps CS-SARI unsafe issues at zero while a no-hold control has "
        "1,649,883 unsafe modeled issues. The ARM64 full-system "
        "path now boots Ubuntu/Linux 6.8.12, reports aarch64, switches from "
        "atomic boot to timing CPU for a native static AArch64 ROI, and "
        "attaches the selected L1D prefetcher. In the tiny ROI, naive "
        "pointer-shaped DMP issues 40 prefetches, none useful, and 30 "
        "translation-faulted recursive attempts; pre-PASB COPPER still permits "
        "5 faulting candidates; PASB-COPPER blocks all 102 pointer-shaped "
        "unproven candidates with zero translation faults. Larger generated full-system ROIs then show CTLW-terminal COPPER at -0.531% ticks on page-permuted pointers and -0.271% on random pointers, with zero recursive translation faults and about 15k unproven candidates blocked per run. A generated full-system AArch64 graph-gather control gives COPPER-CTLW -0.367% ticks, 8,660 blocked unproven candidates, and zero translation faults, while stride wins because the edge array is sequential. LLVM/clang C suites add graph/hash/tree/fake-pointer and GAPBS-inspired graph kernels; COPPER blocks 679 unproven candidates on the first suite, 952 on GAPBS-mini, and 1,340 on a larger 1024-node/degree-8/3-pass GAPBS-inspired run, all with zero translation faults. Bounded checkers pass the full rule, including a richer 11,419-state CEPF/PASB/CTLW authority search, a 24,354-state CLPD search, and a 7,555-state CS-SARI/CLPD/CTLW composition search, while weakened variants fail; full-authority and CEPF-line SVA harnesses pass 10,000 randomized assertion samples each, CTLW witness, CTLW-to-full-authority, CLPD-CTLW, SARI, and CS-SARI revocation RTL pass randomized XSIM checks, and a source-backed coverage matrix passes for ten modeled unsafe classes. The defensible contribution is the committed pointer-provenance authority rule for recursive data-driven prefetch, not a priority claim."
    )

    doc.add_heading("1. Problem and Thesis", level=1)
    add_paragraph(
        doc,
        "DMPs are attractive because ordinary address-stream prefetchers are "
        "often late on linked structures, graph traversals, hash tables, and "
        "arrays of pointers. The security failure is that a DMP can blur data "
        "and addresses beneath the ISA: a passive value in memory can become a "
        "microarchitectural dereference even if software never used that value "
        "as an address."
    )
    add_paragraph(
        doc,
        "COPPER's thesis is narrow: a DMP should not treat address-shapedness "
        "as authority. It should treat committed architectural pointer use as "
        "authority, and every recursive DMP source should remain under that "
        "same rule."
    )
    add_bullets(doc, [
        "Goal: preserve useful DMP behavior for pointer structures while blocking data-at-rest, stale, cross-domain, and permission-invalid dereferences.",
        "Non-goal: preventing all cache side channels or hiding legitimate architectural secret-dependent address generation.",
        "ARM/SoC fit: the rule can be mapped to AArch64-style translation context, ASID/VMID/security state, and AMBA/CHI invalidation events without copying proprietary Arm internals.",
    ])

    doc.add_heading("2. Core Mechanism", level=1)
    add_paragraph(
        doc,
        "COPPER is defined by an invariant, not by a specific table shape:"
    )
    p = doc.add_paragraph()
    add_run(p, "allow_dmp(source_word, value, context) iff ", bold=True)
    add_run(
        p,
        "committed execution previously used that exact source word/value as "
        "an address source, the source has remained clean, context matches, "
        "and target translation/permission succeeds or an exact CTLW record supplies the target line."
    )
    add_paragraph(
        doc,
        "The minimal line-resident version stores one proof bit per pointer "
        "word plus optional domain color. The gem5 prototype uses a bounded "
        "source-line/value proof ledger so clean proofs can survive L1 "
        "replacement, while writes, invalidations, coherence updates, DMA, and "
        "permission failures revoke or block proof."
    )
    add_numbered(doc, [
        "Proof creation: a later committed demand access to the candidate target line turns a recent source word/value into a proof.",
        "Proof use: a DMP candidate may issue only if the source word/value matches an existing proof.",
        "CPTQ: cross-page candidates enter a committed page-translation queue and issue only after valid translation.",
        "RCP: when COPPER issues a prefetch, it records the target line's identity and context. On fill, that line can seed another DMP request only if the ledger already proves the source word/value.",
        "PASB: proofs and carried records include an address-space token so Linux process changes on the same hardware context cannot reuse old authority.",
        "CTLW: recursive cross-page fills use an exact committed virtual-to-physical target-line witness and are terminal until demand-validated.",
    ])

    doc.add_heading("3. Why This Is Not Just Another RTL Block", level=1)
    add_paragraph(
        doc,
        "The RTL gate is small, but the research claim is the authority rule. "
        "Taint tracking marks sensitivity or information flow; COPPER marks a "
        "positive, narrow permission for one microarchitectural consumer. "
        "CHERI, Morello, MTE, and capability/tag systems protect "
        "architectural pointer use or memory safety; COPPER protects a DMP "
        "that can otherwise act on data never architecturally used as an "
        "address. RCP further distinguishes COPPER from ordinary recursive "
        "pointer chasing: a prefetched line carries identity, not permission."
    )

    doc.add_heading("4. Evaluation Artifacts", level=1)
    add_table(
        doc,
        ["Artifact", "What it tests", "Status"],
        [
            ["Python trace model", "Safety invariant under benign/adversarial DMP candidates", "Completed"],
            ["Fuzz validation", "Directed stale, first-use, cross-domain, translation, and permission cases", "500 fuzz trials, 0 failures"],
            ["Vivado RTL", "Gate behavior and rough FPGA cost/timing", "Simulation completed for core/full-authority/CLPD; synthesis reports for subblocks"],
            ["RTL SVA/CTLW/SARI", "Full-authority, CEPF proof flow, CTLW/CLPD consumer gates, and SoC revocation", "SARI/CS-SARI and authority regressions pass randomized XSIM"],
            ["Bounded checkers", "PASB/CTLW/terminal, CEPF/value authority, CLPD, and CS-SARI composition", "Full rules pass; 11,419-state authority, 24,354-state CLPD, and 7,555-state composition searches"],
            ["Coverage matrix", "Unsafe-class-to-evidence audit", "Ten modeled classes; evidence string audit PASS"],
            ["ChampSim", "Stock prefetcher weakness on pointer-shaped traces", "Built and run locally"],
            ["GAPBS", "Graph workload tool readiness", "Built; scale-20 kernels verified"],
            ["GAPBS-backed traces", "Generated graph topology plus COPPER/CLPD proof behavior", "Five graph sizes, kernel sweep, capacity/sensitivity"],
            ["ARM/AArch64 gem5 SE", "Timing caches, MSHRs, L2, DDR3, and COPPER prefetch issue", "COPPER/RCP/CPTQ integrated"],
            ["ARM64 gem5 FS", "Ubuntu/Linux native AArch64 ROIs with L1D prefetcher attached", "PASB, CTLW, graph, C-suite, and small/larger GAPBS-mini runs completed"],
        ],
        [1700, 5600, 2060],
    )

    doc.add_heading("5. Security Trace Results", level=1)
    add_paragraph(
        doc,
        "The trace model intentionally includes benign linked pointer chains, "
        "data-at-rest secret-like values, cross-domain pointer-shaped values, "
        "and rewritten source words."
    )
    add_table(
        doc,
        ["Policy", "Speedup", "Prefetches", "Data-at-rest", "Cross-domain", "Unproven line"],
        [
            ["disabled", "1.000x", "0", "0", "0", "0"],
            ["naive DMP", "3.628x", "4032", "2048", "1076", "2616"],
            ["COPPER-LINE", "2.414x", "1416", "0", "0", "0"],
            ["COPPER-STREAM", "1.641x", "944", "0", "0", "0"],
        ],
        [1400, 1100, 1400, 1500, 1500, 1460],
    )
    add_paragraph(
        doc,
        "Across 30 Monte Carlo seeds, COPPER-LINE maintained zero modeled "
        "unsafe dereferences. In rewrite sensitivity sweeps, safe policies lost "
        "prefetch opportunities as intended but did not degrade into unsafe "
        "data-at-rest activation. A follow-on GAPBS-backed topology trace parses "
        "public .sg graphs and replays edge-scan/BFS streams: naive DMP averages "
        "4.096x but issues 16,384 data-at-rest and 295,463.1 unproven-edge "
        "prefetches, source-only proof still permits 851.3 stale rewritten-slot "
        "prefetches, COPPER-epoch gives 1.770x with zero unsafe counters, and "
        "CLPD gives 1.896x while closing the g12 proof-capacity cliff with 8,192 "
        "line entries; the storage proxy is 54.00 KiB for that CLPD point versus "
        "1696.00 KiB for the edge-exact point that reaches 2.369x. A 4,320-run expanded kernel sensitivity sweep then keeps COPPER unsafe modeled prefetches at 0 while naive DMP produces 81,605,320 unsafe prefetches and source-only provenance still produces 284,488. The CLPD directory RTL now passes 14 directed plus 5,000 "
        "randomized scoreboard checks, including write/fill/invalidate clearing "
        "and collision eviction; a bounded CLPD checker passes 24,354 reachable "
        "states while weakened tag/token/epoch/mask/clear variants fail. A generated "
        "security coverage matrix maps ten unsafe classes to local evidence and "
        "passes evidence string checks; it is an audit, not a complete proof."
    )

    doc.add_heading("6. RTL and Vivado Results", level=1)
    add_table(
        doc,
        ["RTL block", "LUTs", "Registers", "BRAM", "DSP", "WNS @ 10 ns", "Worst path"],
        [
            ["copper_line_provenance_gate", "2063 / 20800", "1024 / 41600", "0", "0", "+8.122 ns", "1.727 ns"],
            ["copper_stream_table_gate", "2528 / 20800", "2209 / 41600", "0", "0", "+0.232 ns", "9.386 ns"],
        ],
        [2500, 1200, 1200, 700, 650, 1200, 1210],
    )
    add_paragraph(
        doc,
        "The line-provenance gate is fast because it is an indexed metadata "
        "check. The stream-table extension is closer to the timing limit due to "
        "direct CAM-like comparisons. This supports presenting COPPER-LINE/RCP "
        "as the core mechanism and COPPER-STREAM as an optional extension. A "
        "new full-authority CEPF/PASB/CTLW gate also passes XSIM with 12 "
        "directed cases and 5,000 named-coverage randomized scoreboard trials "
        "(allowed=956, blocked=3,731, errors=0); fresh synthesis of that gate "
        "is blocked by the current Vivado Tcl initialization failure. A separate "
        "full-authority SVA harness passes 12 directed plus 10,000 randomized "
        "assertion samples and checks allow/source/token/target/terminal/permission "
        "implications directly. A new CEPF-to-line end-to-end SVA harness connects "
        "the commit-epoch bridge to the line gate and passes 12 directed plus "
        "10,000 randomized samples with 2,257 valid commits, 769 proof-to-allow "
        "cases, and 0 errors. A CTLW witness-directory RTL harness passes 10 "
        "directed plus 10,000 randomized samples with exact-hit, line-mismatch, "
        "token-mismatch, remap-clear, TLBI-clear, and collision coverage. A "
        "CTLW-to-full-authority E2E harness passes 12 directed plus 10,000 "
        "randomized samples with exact cross-page allow plus no-witness, "
        "line-mismatch, token-mismatch, remap/TLBI-stale, terminal, permission, "
        "stale-source, and collision blocking at the final gate. A CLPD-CTLW "
        "authority harness adds 18 directed plus 10,000 randomized samples and "
        "checks that compressed source proof and exact target witness must both "
        "be live. SARI adds a SoC revocation queue/hold contract for "
        "DMA/CHI/I/O writes plus remap/TLBI events and passes 10,000 randomized "
        "cycles. CS-SARI refines this with candidate-specific conflict hold: "
        "12 directed plus 10,000 randomized XSIM samples show matching source, "
        "remap, TLBI, queued, and overflow hazards hold while unrelated "
        "revocations safely avoid 1,007 global holds. A composition checker "
        "requires scoped hold, live CLPD source proof, live CTLW target witness, "
        "and source/target clearing to agree; source-only, target-only, "
        "stale-clear, and missing-hold variants all produce counterexamples. "
        "A queue-depth/conflict sweep reports 0 CS-SARI unsafe modeled issues, "
        "1,649,883 no-hold unsafe issues, and 72.06% median hold reduction."
    )

    doc.add_heading("7. ARM gem5 Results", level=1)
    add_table(
        doc,
        ["Workload", "Prefetcher", "Speedup", "Ticks", "Demand MSHR", "PF MSHR", "Carried hits"],
        [
            ["Sequential 8192", "recursive COPPER", "6.76%", "3291910000", "8451", "25260", "25259"],
            ["Sequential 8192", "stride", "98.93%", "1766614000", "522", "32504", "0"],
            ["Page-permuted 8192", "recursive COPPER", "6.76-6.78%", "~3291M", "8451", "~25265", "~25265"],
            ["Page-permuted 8192", "stride", "0.69%", "3490348000", "32776", "250", "0"],
            ["Random 8192", "same-page COPPER", "0.06%", "4116483000", "32828", "198", "0"],
            ["Random 8192", "CPTQ one-hop", "2.74%", "4008993000", "20738", "12288", "0"],
            ["Random 8192", "recursive COPPER", "5.59-5.66%", "~3898M", "8451", "~25167", "~25166"],
            ["Random 8192", "stride", "0.59%", "~4094M", "32776", "250", "0"],
            ["Medium pageperm", "recursive COPPER", "15.06%", "45805000", "263", "818", "817"],
            ["AArch64 pageperm", "recursive COPPER", "6.77%", "3298527000", "8449", "25346", "25345"],
            ["AArch64 random", "recursive COPPER", "5.61%", "3907457000", "8449", "25240", "25239"],
            ["ARM64 FS probe", "no_systemd boot", "env", "250338247164", "-", "-", "-"],
            ["ARM64 FS timing", "stride", "5.279%", "3382963983", "32022", "15443", "7125 useful"],
            ["ARM64 FS timing", "naive DMP", "-0.024%", "3572333757", "38900", "40", "30 xlate faults"],
            ["ARM64 FS timing", "PASB-COPPER", "0.000%", "3571493265", "38899", "0", "102 blocked"],
            ["ARM64 FS pageperm", "CTLW-terminal", "0.531%", "5147204976", "76236", "24229", "0 xlate faults"],
            ["ARM64 FS random", "CTLW-terminal", "0.271%", "5306601753", "75795", "12341", "15259 blocked"],
            ["ARM64 FS graph", "CTLW-terminal", "0.367%", "5253046029", "89234", "49891", "8660 blocked"],
            ["ARM64 FS C suite", "CTLW-terminal", "-0.093%", "4605001389", "84332", "904", "679 blocked"],
            ["ARM64 FS GAPBS-mini", "CTLW-terminal", "0.000%", "3770630928", "47877", "1288", "952 blocked"],
            ["ARM64 FS GAPBS-large", "CTLW-terminal", "-0.208%", "7434010881", "145813", "4953", "1340 blocked"],
            ["Authority checker", "CEPF/PASB/CTLW", "PASS", "11419 states", "-", "-", "bugs fail"],
        ],
        [1650, 1700, 980, 1300, 1240, 1240, 1250],
    )
    add_paragraph(
        doc,
        "The key gem5 result is not a lower raw D-cache miss count; it is a "
        "conversion of demand-visible MSHR misses into prefetch-origin MSHR "
        "misses. On ARM32 random seed 1, demand MSHR misses fall from 33,026 to "
        "8,451 while recursive COPPER issues 25,166 proof-gated prefetches. "
        "On AArch64 random, demand MSHR misses fall from 33,024 to 8,449. "
        "Same-page COPPER barely helps random lists, CPTQ recovers cross-page "
        "one-hop issue, and RCP provides the strongest result by allowing safe "
        "recursive runahead. The ARM64 full-system path ran Linux 6.8.12, "
        "mounted /dev/vda2, loaded gem5_bridge, reported aarch64, completed a "
        "470.7M-instruction probe, and ran a bracketed native AArch64 timing "
        "ROI with the L1D prefetcher attached. The first key lesson was PASB: naive "
        "DMP produced 30 translation-faulted recursive attempts, pre-PASB "
        "COPPER left 5, and PASB-COPPER left 0 while blocking 102 unproven "
        "pointer-shaped candidates. The larger generated pointer ROIs then exposed CTLW: exact target-line witnesses plus terminal witness-derived fills removed PASB-only translation faults while preserving small positive timing movement. The generated graph-gather run adds a CSR-like full-system control, while the clang-authored C and GAPBS-inspired mini-suites show COPPER's authority block and zero translation faults on compiler-generated code; the larger GAPBS-inspired rerun blocks 1,340 unproven candidates, records 50,737 terminal stops, removes naive's CTLW misses/unavailable translations, and gives -0.208% ticks versus no prefetch. The GAPBS-backed topology replay is not full-system GAPBS, but it now covers five generated graph sizes, four kernel access shapes, and a 4,320-run sensitivity sweep; it exposed a proof-capacity cliff and gave CLPD as a measured fix. CS-SARI adds a workload-derived revocation proxy over those graph topologies: 6.34M raw candidates per scenario, 58-97% hold reduction versus global hold, zero modeled unsafe issues, and 269,879 authorized opportunities recovered in aggregate; a later queue-depth/conflict sweep broadens this to 20 configurations with zero CS-SARI unsafe issues and 1,649,883 no-hold unsafe issues. CLPD and CS-SARI now have fresh RTL simulation/proxy evidence, but not fresh synthesis because Vivado TclStore/app loading fails before read_verilog/synth_design. These graph-kernel controls are neutral or negative for speed, which the paper should report honestly. The richer bounded checker passes the full authority rule over 11,419 reachable states to depth 12 and gives short counterexamples for stale backend proof, missed source invalidation, address-space reuse, page-level witnesses, stale remap witnesses, and terminal-fill recursion. The full-authority and CEPF-line SVA harnesses add gate-level and multi-cycle property evidence, CTLW/CLPD-CTLW RTL add consumer-gate evidence, SARI/CS-SARI add SoC revocation hold evidence, and the coverage matrix adds a ten-class source-backed audit, while leaving residual production risks explicit."
    )

    doc.add_heading("8. Prior Art Position", level=1)
    add_table(
        doc,
        ["Prior art", "Similarity", "Difference from COPPER"],
        [
            ["Augury / GoFetch", "Very high threat overlap", "Attack/reverse-engineering work; no committed source-word authority rule."],
            ["SplittingSecrets", "High DMP defense overlap", "Compiler transforms secrets so they do not look like addresses; COPPER changes hardware DMP authority."],
            ["PreFence / DIT / DOIT disable", "Medium defense overlap", "Coarse prefetcher disable; COPPER is fine-grained and preserves safe DMP activity."],
            ["Pointer-chase / indirect prefetchers / ICP", "High performance overlap", "Performance mechanisms without committed source-word proof for DMP dereference."],
            ["Taint, CHERI/Morello, MTE", "Medium metadata/provenance overlap", "Architectural safety or information-flow metadata, not DMP-specific positive authority."],
        ],
        [1900, 2300, 5160],
    )
    add_paragraph(
        doc,
        "After public search, the defensible claim is not absolute novelty. It "
        "is: to the best of public knowledge, no public work was found that "
        "uses committed clean pointer-source provenance as a hard authority "
        "rule for recursive DMP dereference. Broad novelty risk remains 4/10; "
        "the narrow RCP/CPTQ claim is currently estimated at 3/10."
    )

    doc.add_heading("9. Reviewer Risk Assessment", level=1)
    add_bullets(doc, [
        "Gem5 scope: validated ARM-system summaries and retained raw-run manifests provide full-system AArch64 evidence for rows marked PASS, but the package must not present this as a complete raw-run GAP/SPEC/crypto campaign.",
        "Backend proof path: a production core needs measured load-store-queue or dependency-tag integration to identify committed pointer-source words.",
        "First-use loss: COPPER intentionally cannot prefetch through a source word before committed proof exists.",
        "Industry secrecy: commercial DMP safety rules may exist without public documentation, so claims must remain public-knowledge scoped.",
        "Traffic and pollution: recursive COPPER adds prefetch-origin MSHR pressure; the current result is positive, but broader workloads could expose bandwidth costs.",
    ])

    doc.add_heading("10. Scores and Verdict", level=1)
    add_table(
        doc,
        ["Dimension", "Score", "Reason"],
        [
            ["Novelty risk", "3/10 narrow, 4/10 broad", "Exact RCP/CPTQ authority rule not found publicly, but adjacent fields are crowded."],
            ["Feasibility", "8/10", "Trace model, RTL simulation/synthesis, gem5 cache-model integration, and ARM64 full-system timing ROI all run locally."],
            ["Measurability", "8/10", "Unsafe prefetch counts, demand MSHR misses, translated prefetches, and timing are measurable."],
            ["Hardware cost", "7.5/10", "Core gate is cheap; CLPD storage proxy is strong; backend proof tracking and fresh CLPD synthesis remain open."],
            ["Paper strength", "Evidence-bounded regular submission candidate", "Strong mechanism, full-system CTLW timing plus witness, consumer-gate, CLPD-CTLW, SARI/CS-SARI RTL, graph/C/small-larger GAPBS-mini controls, expanded kernel sensitivity, CLPD capacity/storage/checker evidence, composition and sensitivity checks, SVA evidence, security audit, and richer bounded checker; broader raw-run coverage remains open."],
        ],
        [1800, 1500, 6060],
    )
    add_paragraph(
        doc,
        "Final verdict: evidence-bounded regular-conference or artifact-track candidate, "
        "not an acceptance guarantee. The remaining "
        "weaknesses are broader raw-run workload coverage and production-grade integration, "
        "not the basic ARM64/Linux execution, graph/C controls, GAPBS-backed topology replay, prefetcher-attachment, bounded invariant check, CLPD capacity fix, or translation-witness consumer path."
    )

    doc.add_heading("11. Remaining Evidence Gates", level=1)
    add_numbered(doc, [
        "Move from GAPBS-backed topology replay and GAPBS-inspired mini-suites to official GAPBS full-system AArch64 gem5 runs.",
        "Run real full-system AArch64 workloads that can stress COPPER-CTLW beyond generated/freestanding binaries.",
        "Extend formal/SVA properties through CLPD plus CTLW, real remap/TLBI paths, and coherence/DMA invalidation.",
        "Prototype a backend source-word tag path or an RTL-level commit/proof interface, not only the cache-side gate.",
        "Measure traffic, pollution, bandwidth, and energy proxies under larger cache and memory configurations.",
    ])

    doc.add_heading("References", level=1)
    refs = [
        "Augury: Using Data Memory-Dependent Prefetchers to Leak Data at Rest. https://www.prefetchers.info/augury.pdf",
        "GoFetch: Breaking Constant-Time Cryptographic Implementations Using Data Memory-Dependent Prefetchers. https://gofetch.fail/",
        "SplittingSecrets: A Compiler-Based Defense for Preventing Data Memory-Dependent Prefetcher Side-Channels. https://arxiv.org/abs/2601.12270",
        "PreFence: A Scheduling-Aware Defense Against Prefetching-Based Side-Channel Attacks. https://arxiv.org/abs/2410.00452",
        "ICP: Exploiting Instruction Correlation for Prefetching Irregular Memory Accesses. https://arxiv.org/abs/2605.15645",
        "gem5, ChampSim, and GAP Benchmark Suite public project artifacts.",
    ]
    add_bullets(doc, refs)

    add_footer(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
